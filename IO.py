from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import subprocess
import datetime
import os
import pandas as pd

from Wählbär import MetaBlock, Schedule, Block, Unit, Allocation, SLOTS_PER_DAY, DAYS
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import xlsxwriter as xls

class FORMAT:
    RED = "\033[31m"
    YELLOW = "\033[33m"
    GREEN = "\033[32m"
    BOLD = "\033[1m"
    RESET = "\033[0m"


   


def slot_to_table_idx(slot):
    if type(slot) == str:
        idx = Schedule.to_idx(slot)
    elif type(slot) == tuple:
        idx = slot
    else:
        print(f"ERROR: invalid type {type(slot)} expected 'tuple' or 'str'")
        
    row = idx[1] +1
    if idx[0] <= 6:
        col = idx[0] 
        return 0, row, col
    else:
        col = idx[0] -6 
        return 2, row, col


def export_to_pdf(unit):

    doc = Document("templates/template_unit.docx")
    for day in range(DAYS):
        for time in range(SLOTS_PER_DAY):
            slot = Schedule.idx2str(day, time)
            block = unit.schedule[slot]
            placeholder_ID = "{"+ f"{slot}_id"+ "}"
            placeholder_fullname = "{"+ f"{slot}_fullname"+ "}"
            if len(block) == 1:
                block = block[0]
                id_short = block.ID.split("_")[0] # remove ON-11_XYZ -> ON-11
                if id_short in ["ON-01", "ON-40", "ON-41", "ON-42", "ON-43"]:
                    id_short += "*"
                if id_short[:3] == "OFF":
                    id_short +="**"
                replace_text_in_document(doc, placeholder_ID, id_short)
                replace_text_in_document(doc, placeholder_fullname, block.data["fullname"])
            elif len(block) == 0:
                replace_text_in_document(doc, placeholder_ID, "")
                replace_text_in_document(doc, placeholder_fullname, "")
            elif len(block) > 1:
                print(f"ERROR: more than one block in slot {slot} of unit {unit.ID}")
                for i, block in enumerate(block):
                    print("  -  ", block.ID)
    group_map = {
        "pf": "Pfadis",
        "pi": "Pios",
        "wo": "Wölfe"
    }

    for placeholder, value in { # TODO
    "{name}": unit.fullname,
    "{date}": datetime.datetime.now().date().strftime("%d.%m.%Y"),
    "{ID}": str(unit.ID),
    "{group}": group_map.get(unit.group, "Unbekannt"),
    "{B1_id}": "TEST"}.items():
        replace_text_in_document(doc, placeholder, value)

    color_map = {
        "anlass": "#000000",
        "ausflug": "#f2c966",
        "wanderung": "#00b48f",
        "sportaktivitat": "#c6464a",
        "programmflache": "#e87928",
        "wald": "#e87928",
        "nacht": "#e87928",
        "wasser": "#608ee4",
        "flussbaden": "#608ee4",
        "si-mo": "#608ee4",
        "dusche": "#608ee4",
        "workshop": "#4f2c1d",
    }




    for block in unit.schedule.get_list(with_slot=True):
        tab, row, col = slot_to_table_idx(block["slot"])
        table = doc.tables[tab]

        # Access the cell (e.g., first row, first column)
        cell = table.cell(row, col)

        # Set the background color of the cell
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_map.get(block["element"].data["cat"], "#808080"))  # default white color
        cell._tc.get_or_add_tcPr().append(shading_elm)



    doc.save(f"exports/{unit.ID}.docx")

    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", "exports",
        f"exports/{unit.ID}.docx"
    ])


def replace_text_in_paragraph(paragraph, placeholder, replacement):
    for run in paragraph.runs:
        if placeholder in run.text:
            # Split the text into parts before, after, and the placeholder
            split = run.text.split(placeholder)
            pre = split[0]  # Keep the text before the placeholder
            post = split[1]
            
            run.text = pre + replacement + post  # Keep the text before the placeholder
            # Add the replacement text with the same style
            # for part in parts[1:]:
            #     run = paragraph.add_run(replacement + part)
            #     # Copy the style from the previous run
            #     run.bold = paragraph.runs[-2].bold
            #     run.italic = paragraph.runs[-2].italic
            #     run.underline = paragraph.runs[-2].underline
            #     run.font.name = paragraph.runs[-2].font.name
            #     run.font.size = paragraph.runs[-2].font.size
            #     run.font.color.rgb = paragraph.runs[-2].font.color.rgb

def replace_text_in_document(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, placeholder, replacement)
    
    for section in doc.sections:
        # Header
        if section.header is not None:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, placeholder, replacement)
        # Footer
        if section.footer is not None:
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, placeholder, replacement)


def load_unitlist(allocation, path="data", filename="Antworten Buchungstool.xlsx", print_enabled=False):
    full_labels = pd.read_excel(os.path.join(path, filename), sheet_name="Formularantworten 1", header=1).columns.tolist()
    df = pd.read_excel(os.path.join(path, filename), sheet_name="Formularantworten 1", header=2)
    non_empty_count = df.iloc[:, 0].count() +1
    df = df.head(non_empty_count)

    df.dropna(subset=["ID_all_int"], inplace=True)

    for ic, col in enumerate(df.columns):
        new = col
        display_full_label = full_labels[ic].replace("\n", " ")
        if len(display_full_label) > 80:
            display_full_label = display_full_label[:77] + "..."
        if col.startswith("Unnamed") and full_labels[ic].startswith("Unnamed"):
            if print_enabled: print(f"\033[31m{new:>20}\033[0m: {display_full_label} (will get removed)")
        elif col.startswith("Unnamed") and not full_labels[ic].startswith("Unnamed"):
            try:
                ID = full_labels[ic].split("[")[2].split("]")[0]  
                group = full_labels[ic].split(" - ")[1].split(" ")[0][:2].lower()
                if group == "wö":
                    group = "wo"
                new = f"{ID}_{group}_m3"
                df.rename(columns={col: new}, inplace=True)
                if print_enabled: print(f"\033[32m{new:>20}\033[0m: {display_full_label}")
            except:
                if print_enabled: print(f"\033[31m{new:>20}\033[0m: '{display_full_label}' will get removed for no parsing")
                new = full_labels[ic]
        elif not col.startswith("Unnamed") and not full_labels[ic].startswith("Unnamed"):
            if print_enabled: print(f"\033[32m{new:>20}\033[0m: {display_full_label}")
        else:
            print(f"\033[31mERROR Label with no fullname \033[0m: '{display_full_label}' will get removed")
            # new = full_labels[ic]

    
    for col in df.columns:
        if col.startswith("Unnamed"):
            df.drop(columns=[col], inplace=True)

    if print_enabled:
        print("\033[32mFinal columns:\033[0m")
        for col in df.columns:
            print(f"{col:>20}") 

    group_df = []
    for group in ["Pios", "Pfadis", "Wölfe"]:
        df_group = df[df["group_all_tx"] == group].reset_index(drop=True)
        df_group = df_group[[c for c in df.columns if f"_{group[:2].lower().replace('ö', 'o')}_" in c or "_all" in c]]

        for ic, col in enumerate(df_group.columns):
            # print(f"{'_'.join(col.split('_')[:-2]):>20}") 
            if col.endswith("_m3") or col.endswith("_02") or col.endswith("_03")  or col.endswith("_05") or col.endswith("_int"):
                df_group.fillna({col:-1}, inplace=True)
                df_group = df_group.astype({col:"int32"})
            if col.endswith("_tx"):
                df_group = df_group.astype({col:"string"})
            if col.endswith("_jn"):
                df_group[col] = df_group[col].map({"Ja": True, "Nein": False})
                df_group = df_group.astype({col:"bool"})
            df_group.rename(columns={col: '_'.join(col.split('_')[:-2])}, inplace=True)
        group_df.append(df_group)
        
        
        df_group["OFF-26"] = df_group["AUX-HB"]
        df_group["OFF-27"] = df_group["AUX-HB"]
        df_group["OFF-28"] = df_group["AUX-FR"]
        df_group["OFF-29"] = df_group["AUX-FR"]
        df_group.drop(columns=["AUX-HB", "AUX-FR"], inplace=True)

    df_pi, df_pf, df_wo = group_df

    unclassified = [c for c in df.columns if ("_pi_" not in c) and ("_pf_" not in c) and ("_wo_" not in c) and ("_all_" not in c) and ("ID" not in c)]
    if print_enabled:
        if len(unclassified) >0:
            print("\033[31mUnclassified columns (not assigned to any group):\033[0m")
            for col in unclassified:
                print(f"{col:>20}")
        else:
            print("\033[32mAll columns classified into groups.\033[0m")


    if print_enabled:
        for df_, name in zip([df_pi, df_pf, df_wo], ["Pios", "Pfadis", "Wölfe"]):
            print(f"\033[32m{name} Final columns:\033[0m")
            for col in df_.columns:
                print(f"{col:>20}: {df_.dtypes[col]}")
       
    tn_numbers = pd.read_excel(os.path.join(path, "EH_Übersicht_Einheiten.xlsx"), sheet_name="Übersicht_Einheiten", header=1)
    non_empty_count = tn_numbers.iloc[:, 0].count() +1
    tn_numbers = tn_numbers.head(non_empty_count)
    tn_numbers = tn_numbers[["ID", "Verantwortliche Abteilung", "Teilnehmende", "Betreuungsperson", "Datum"]]
    tn_numbers = tn_numbers.astype({"ID": "string"}).set_index("ID")

    for df_ in [df_pi, df_pf, df_wo]:
        for i in range(df_.shape[0]):
            ID = str(int(df_.loc[i, "ID"]))
            data = {col: df_.loc[i, col] for col in df_.columns[1:]}
            data["group"] = data["group"][:2]
            if tn_numbers.index.str.contains(ID).sum() == 0:
                print(f"\033[33mWARNING: could not find unit {ID} in TN\033[0m")
                data["n_people"] = -1
                data["fullname"] = "UNKNOWN"
            else:  
                data["n_people"] = tn_numbers.loc[ID, "Teilnehmende"]
                data["fullname"] = tn_numbers.loc[ID, "Verantwortliche Abteilung"]
            if tn_numbers.loc[ID, "Datum"] == "12.-25. Juli 2026":
                data["present_on"] = [e-12 for e in range(12, 25+1)] # -12 to convert to wählbär day (12.7. is Day 0)
            if tn_numbers.loc[ID, "Datum"] == "13.-18. Juli 2026":
                data["present_on"] = [e-12 for e in range(13, 18)]
            if tn_numbers.loc[ID, "Datum"] == "20.-25. Juli 2026":
                data["present_on"] = [e-12 for e in range(20, 25+1)]

            allocation.append_unit(
                Unit(
                    ID,
                    data       
            )
        )

    for row in tn_numbers.itertuples():
        ID = row.Index
        if all([ID != unit.ID for unit in allocation.UNITS]):
            print(f"\033[33mWARNING: Unit {ID} is not in booking tool responses ({row.Betreuungsperson}).\033[0m")
    input("Press Enter to continue...")
    print(f"Loaded {len(allocation.UNITS)} units.")
    
    
def load_blocklist(allocation, path="data", filename="PRG_Blockliste.xlsx"):

    # TODO: load unit list
    df = pd.read_excel(os.path.join(path, filename), sheet_name="On-Site Buchbar")
    non_empty_count = df.iloc[:, 0].count() +1
    df = df.head(non_empty_count)
    df = df[[
        'Block Nr.',
        'Status',
        'Kategorie', 
        'Titel',
        'Programmstruktur', 
        'Dauer', 
        'Blockart J+S', 
        'Stufe', 
        'Gruppengrösse',
        "Hard limit", 
        'Partizipation',
        'max. Anzahl Durchführungen (wie viele Einheiten können diesen Block besuchen?)',
        'geschätzte Anzahl Durchführungen',
        'Durchführingszeiten (Start)',

        'Verteilungsprio'
    ]]

    df.columns = [
        'ID',
        'state',
        'cat', 
        'fullname',
        'betr_unbetr', 
        'dauer', 
        'blockart_J_S', 
        'stufen', 
        'gruppengroesse',
        "hard_limit", 
        'mix_units',
        'max_durchfuhrungen',
        'est_durchfuhrungen',
        'start_times',

        'verteilungsprio'
    ]

    df_offsite = pd.read_excel(os.path.join(path, filename), sheet_name="Off Site & Wasser")
    non_empty_count = df_offsite.iloc[:, 0].count() + 1
    df_offsite = df_offsite.head(non_empty_count)
    df_offsite = df_offsite[[
        'Block Nr.',
        'Status',
        'Kategorie', 
        'Titel',
        'Programmstruktur', 
        'Dauer', 
        'Blockart J+S', 
        'Stufe', 
        'Gruppengrösse', 
        "Hard limit",
        'Partizipation',
        'max. Anzahl Durchführungen (wie viele Einheiten können diesen Block besuchen?)',
        'geschätzte Anzahl Durchführungen',
        'Durchführingszeiten (Start)',

        "Verteilungsprio"

    ]]

    df_offsite.columns = [
        'ID',
        'state',
        'cat', 
        'fullname',
        'betr_unbetr', 
        'dauer', 
        'blockart_J_S', 
        'stufen', 
        'gruppengroesse',
        "hard_limit", 
        'mix_units',
        'max_durchfuhrungen',
        'est_durchfuhrungen',
        'start_times',

        'verteilungsprio'
    ]

    df = pd.concat([df, df_offsite], ignore_index=True)
    df = df.dropna(subset=["ID"])
    for bd in df.itertuples():
        
        length = 1
        on_times = [0, 1, 2, 3, 4]
        if bd.dauer == "4 h": 
            length = 2
            on_times = [1]
        if bd.dauer == "6 h": 
            length = 3
            on_times = [0]
        if bd.dauer == "8 h":
            length = 3
            on_times = [0]
        if bd.dauer == "2 Tage":
            length = 8
            on_times = [0]
        
        slots = []
        if not pd.isna(bd.start_times):
            start_array = bd.start_times.split(", ")
            day = [int(e[:2]) - 12 for e in start_array]
            time = [ord(e[2]) - 65 for e in start_array]
            for i in range(len(day)):
                slots.append(Schedule.idx2str(day[i], time[i]))

        group = [e[:2].lower().replace('ö', 'o') for e in bd.stufen.split(', ')]

        cat = bd.cat.lower().replace("ä", "a")
        if cat == "zweitageswanderung":
            cat = "wanderung"

        tags = set()
        if "2 Einheiten" in bd.fullname:
            tags.add("2units")
        mix_units = False
        if not pd.isna(bd.mix_units):
            if "2/3 Einheiten zusammen" in bd.mix_units or "ganzes KALA" in bd.mix_units:
                mix_units = True

        allocation.append_block(
            Block(
                bd.ID,
                {   
                    "ID": bd.ID,
                    "fullname": bd.fullname,
                    "space": bd.gruppengroesse,
                    "hard_limit": bd.hard_limit,
                    "js_type": bd.blockart_J_S,
                    "cat":  cat,
                    "group": group,
                    "length": length,
                    "on_days": [1, 2, 3, 4, 5, 6, 8, 9, 10, 11, 12],
                    "on_times": on_times,
                    "on_slots": slots,
                    "tags": tags,
                    "state": bd.state,

                    "verteilungsprio": bd.verteilungsprio,
                    "mix_units": mix_units
                }
            )
        )
    
    allocation.find_block_cats()
    print(f"Loaded {len(allocation.BLOCKS)} blocks.")

 
def plot_block(ax, slot, block):
    day, slot = Schedule.to_idx(slot)
    x = day
    y = SLOTS_PER_DAY - slot 
    ax.add_patch(patches.Rectangle((x, y), 1, -block.requirements["length"], linewidth=1, edgecolor='none', facecolor='teal'))
    ax.text(x + 0.1, y-0.3, block.ID, fontweight="bold")
    ax.text(x + 0.1, y-0.5, "ORT")


def print_schedule(thing, save=False):
    if type(thing) == Block or type(thing) == Unit:
        schedule = thing.schedule
    elif type(thing) == Schedule:
        schedule = thing
    else:
        print("ERROR: unknown schedule type. expected 'Block', 'Unit' or 'Schedule'")

    fig = plt.figure(figsize=(16, 5))
    ax = plt.subplot(111)
    for idd, day in enumerate(schedule.calendar):
        for iss, slot in enumerate(day):
            if len(slot) == 1:
                block = slot[0]
                plot_block(ax, (idd, iss), block)
            if len(slot) > 1:
                print("ERROR: MORE THAN ONE BLOCK PER SLOT")
    ax.set_xlim(-1, 15)
    ax.set_ylim(-1, 5)

    ax.set_xticks([i + 0.5 for i in range(14)])
    # TODO: Labels for days
    
    if save:
        plt.savefig(save)
    else:
        plt.show()

def print_block_geilheit(a):
    counter = {}
    for block in a.BLOCKS:
        if "AUX" not in block.ID:
            counter[block.ID] = 0

    for unit in a.UNITS:
        for ID, value in unit.prios.items():
            if len(ID.split("-")) == 2 and "AUX" not in ID:
                counter[ID] += value
    
    group_conter = {"pf": 0, "wo": 0, "pi": 0, "ro": 0}
    for unit in a.UNITS:
        group_conter[unit.group] +=1
    
    for block in a.BLOCKS:
        if not "AUX" in block.ID:
            allowed_units = sum([group_conter[gr] for gr in block.data["group"]])
            counter[block.ID] /= allowed_units    

    x = range(len(counter.keys()))
    y = counter.values()
    labels = counter.keys()
    plt.figure(figsize=(16, 5))
    plt.bar(x, y, color="#00b48f")
    plt.xticks(x, labels, rotation="vertical")
    plt.ylabel("Durchschnitt Prio Punkte aller Einheiten")
    plt.tight_layout()
    plt.show()
       

def slot_to_xlsx_cell(slot):
    day, time = ord(slot[0]) - 65, int(slot[1])
    return f"{chr(day+1+65) + str(time+1 +2)}"


def write_to_xlsx(allocation, fname="allocation.xlsx", path="saves"):
    workbook = xls.Workbook(os.path.join(path,fname))

    merge_format = workbook.add_format(
        {
            "bold": 1,
            "border": 0,
            "align": "center",
            "valign": "vcenter",
            # "fg_color": "yellow",
        }
    )

    group_colors = {
        "wo": "#0db3bb",
        "pf": "#9c8566",
        "pi": "#c51f1f"
    }
    green_format = workbook.add_format({"font_color": "#00b48f"})
    red_format = workbook.add_format({"font_color": "#f88589"})
    blue_format = workbook.add_format({"font_color": "#608ee4"})
    allocation.UNITS = sorted(allocation.UNITS, key=lambda e: e.ID) 
    for iu, unit in enumerate(allocation.UNITS):
        # unit = allocation.UNITS[0]
        # print(f"Writing unit {unit.ID} to xlsx...")
        worksheet = workbook.add_worksheet(unit.ID)
        worksheet.set_tab_color(group_colors[unit.group])
    
        worksheet.merge_range("B1:O1", f"{unit.ID}: {unit.fullname} ({unit.group})", merge_format)
        for i in range(1, SLOTS_PER_DAY +1):
            worksheet.write(f"A{i+2}", f"slot {i}")
        for i in range(DAYS):
            worksheet.write(f"{chr(i+1+65)}2", f"{i+12}.07.")

        for blocks in unit.schedule.get_time_list():
            worksheet.write(slot_to_xlsx_cell(blocks["slot"]), "/".join(blocks["elements"]))

        row = 10; col = 0

        worksheet.write(f"{chr(col+65)}{row}", "info"); row+=1; col+=1
        worksheet.write(f"{chr(col+65)}{row}", f"n_people: {unit.n_people}"); col+=1
        
        row +=1; col = 0
        worksheet.write(f"{chr(col+65)}{row}", "general"); row+=1; col+=1
        for key, value in unit.general.items():
            if col > 8:
                col = 1
                row +=1
            worksheet.write(f"{chr(col+65)}{row}", f"{key}:"); col+=1
            worksheet.write(f"{chr(col+65)}{row}", f"{value}"); col+=1
        
        row+=1; col = 0
        for cat, prios in unit.prios_sorted.items():
            worksheet.write(f"A{row}", cat)
            row +=1
            col = 1
            for prio in prios:
                if col > 4:
                    col = 1
                    row +=1
                if prio['value'] >= 2 and unit.has_block(prio['ID']):
                    worksheet.write(f"{chr(col+65)}{row}", f"{prio['ID']}: {prio['value']}", green_format)
                elif prio['value'] >=2 and not unit.has_block(prio['ID']):
                    worksheet.write(f"{chr(col+65)}{row}", f"{prio['ID']}: {prio['value']}", red_format)
                elif prio['value'] <2 and unit.has_block(prio['ID']):
                    worksheet.write(f"{chr(col+65)}{row}", f"{prio['ID']}: {prio['value']}", blue_format)
                else:
                    worksheet.write(f"{chr(col+65)}{row}", f"{prio['ID']}: {prio['value']}")
                col +=1
            row+= 1
    # worksheet = workbook.add_worksheet("Freie Blöcke")
    # worksheet.merge_range("B1:O1", "Freie Blöcke", merge_format)
    # for i in range(1, SLOTS_PER_DAY +1):
    #     worksheet.write(f"A{i+2}", f"slot {i}")
    # for i in range(DAYS):
    #     worksheet.write(f"{chr(i+1+65)}2", f"day {i+1}")

    # for idd in range(DAYS):
    #     for itt in range(SLOTS_PER_DAY):
    #         cell = chr(idd + 1+65) + str(itt +3)
    #         string = "=CONCATENATE(CONCATENATE("
    #         for ib, block in enumerate(allocation.BLOCKS):
    #             string += f'IF(${block.ID}.{cell} ="";CONCATENATE(${block.ID}.B1; CHAR(10));"");'
    #             if ib == 63:
    #                 string = string[:-1] + "); CONCATENATE("
            
    #         string = string[:-1]+"))"
    #         worksheet.write(cell, string)   
    cat_colors = {
        "flussbaden": "#98b3e4",
        "wasser": "#608ee4",
        "si-mo": "#2052b0",
        "sportaktivitat": "#c6464a",
        "programmflache": "#f88589",
        "ausflug": "#f2c966",
        "workshop": "#e87928",
        "wanderung": "#00b48f",
        "nacht": "#094F41", 
        "wald": "#49ebca",
        "dusche": "#673B80",
        "amtli": "#673B80",
        "AUX": "#808080",
        "anlass": "#673B80"
    }
    
    for ib, block in enumerate(allocation.BLOCKS):
        if isinstance(block, MetaBlock):
            continue
        # unit = allocation.UNITS[0]

        worksheet = workbook.add_worksheet(block.ID)
        if block.is_active:
            worksheet.set_tab_color(cat_colors[block.data["cat"]])
        else:
            worksheet.set_tab_color("#808080")
        worksheet.merge_range("B1:O1", f"{block.ID}: {block.data['fullname']} ({block.data['cat']})", merge_format)
        for i in range(1, SLOTS_PER_DAY +1):
            worksheet.write(f"A{i+2}", f"slot {i}")
        for i in range(DAYS):
            worksheet.write(f"{chr(i+1+65)}2", f"{i+12}.07.")
        

        for units in block.schedule.get_time_list():
            worksheet.write(slot_to_xlsx_cell(units["slot"]), "/".join(units["elements"]))

        col = 0; row = 10
        worksheet.write(f"{chr(col+65)}{row}", "info"); row+=1; col+=1
        for key, value in block.data.items():
            worksheet.write(f"{chr(col+65)}{row}", f"{key}:" )
            worksheet.write(f"{chr(col+65+1)}{row}", f"{value}")
            row+=1

    workbook.close()

